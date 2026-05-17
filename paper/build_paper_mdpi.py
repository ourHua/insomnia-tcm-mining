"""build_paper_mdpi.py — Insomnia_TCM_Mining_Manuscript_MDPI.docx

Uses the MDPI Applied Sciences template (applsci-template.dot, content-type
patched to docx) as the type-setting base. The template's native styles are
applied verbatim:

    MDPI_1.1_article_type       Article / Review / Communication tag
    MDPI_1.2_title              Title (Palatino Linotype 18 bold)
    MDPI_1.3_authornames        Author block
    MDPI_1.6_affiliation        Affiliations (8 pt, hanging indent)
    MDPI_1.7_abstract           Abstract (12 pt justified)
    MDPI_1.8_keywords           Keywords
    MDPI_1.9_line               Page-rule separator before Introduction
    MDPI_2.1_heading1           Heading 1 (12 pt bold)
    MDPI_2.2_heading2           Heading 2 (italic)
    MDPI_2.3_heading3           Heading 3
    MDPI_3.1_text               Body text, first-line indent 21.25 pt
    MDPI_3.2_text_no_indent     Body text without indent
    MDPI_3.3_text_space_after   Body text + space-after 12 pt
    MDPI_3.4_text_space_before  Body text + space-before 12 pt
    MDPI_3.7_itemize            Numbered list item
    MDPI_3.8_bullet             Bullet list item
    MDPI_4.1_table_caption      Table caption (9 pt)
    MDPI_4.3_table_footer       Table footer (9 pt)
    MDPI_5.1_figure_caption     Figure caption (9 pt)
    MDPI_5.2_figure             Centred figure paragraph
    MDPI_6.2_back_matter        Author contributions / DAS / etc. (9 pt)
    MDPI_8.1_references         References (9 pt)
"""
from __future__ import annotations

import copy
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path("/home/claude/template_work/applsci-template-fixed.docx")
OUT = ROOT / "paper" / "Insomnia_TCM_Mining_Manuscript_MDPI.docx"
FIG = ROOT / "outputs" / "figures"


# ────────────────────────────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────────────────────────────
def _clear_body(doc):
    """Strip every paragraph/table from the document body but keep sectPr.

    Also removes the section-level <w:bidi/> flag, which the MDPI template
    ships with set to ON and which causes table columns and bullet markers
    to be rendered right-to-left in LibreOffice and some Word versions.
    """
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sectPr:
            body.remove(child)
    # Remove section-level bidi (LTR documents only)
    if sectPr is not None:
        bidi = sectPr.find(qn("w:bidi"))
        if bidi is not None:
            sectPr.remove(bidi)


def _para(doc, text, style):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    return p


def _para_runs(doc, segments, style):
    """Create a paragraph with mixed bold/italic segments."""
    p = doc.add_paragraph(style=style)
    for seg in segments:
        if isinstance(seg, str):
            p.add_run(seg)
        else:
            r = p.add_run(seg["t"])
            if seg.get("b"):
                r.bold = True
            if seg.get("i"):
                r.italic = True
    return p


def _bullet(doc, text):
    return _para(doc, text, "MDPI_3.8_bullet")


def _figure(doc, png_path: Path, caption_text: str, width_in: float = 5.6):
    fp = doc.add_paragraph(style="MDPI_5.2_figure")
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if png_path.exists():
        fp.add_run().add_picture(str(png_path), width=Inches(width_in))
    else:
        fp.add_run(f"[Figure missing: {png_path.name}]").italic = True
    _para(doc, caption_text, "MDPI_5.1_figure_caption")


def _set_table_borders(table):
    """Apply MDPI-style three-line borders: top + bottom of header, bottom of body."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        if edge in ("top", "bottom"):
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")
            el.set(qn("w:color"), "000000")
        elif edge == "insideH":
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    # Replace existing borders
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(borders)


def _table(doc, rows, col_widths_cm=None):
    n_rows = len(rows)
    n_cols = len(rows[0])
    t = doc.add_table(rows=n_rows, cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths_cm:
        for col_idx, w in enumerate(col_widths_cm):
            for row in t.rows:
                row.cells[col_idx].width = Cm(w)
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = t.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.style = doc.styles["MDPI_4.2_table_body"]
            run = p.add_run(str(cell_text))
            if r_idx == 0:
                run.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_table_borders(t)
    return t


# ────────────────────────────────────────────────────────────────────────
# Build
# ────────────────────────────────────────────────────────────────────────
def build():
    doc = Document(str(TEMPLATE))
    _clear_body(doc)

    # ── Front matter ────────────────────────────────────────────────
    _para(doc, "Article", "MDPI_1.1_article_type")
    _para(doc,
          "Association Rule Mining and Complex Network Analysis of "
          "Insomnia Traditional Chinese Medicine Prescriptions: A "
          "Reproducible Computational Pipeline",
          "MDPI_1.2_title")

    # Author line
    p = doc.add_paragraph(style="MDPI_1.3_authornames")
    p.add_run("Junhua Li ")
    p.add_run("1").font.superscript = True
    p.add_run(" and Gilja So ")
    p.add_run("1,*").font.superscript = True

    # Affiliations
    p1 = doc.add_paragraph(style="MDPI_1.6_affiliation")
    p1.add_run("1").font.superscript = True
    p1.add_run("\tAffiliation to be inserted, Republic of Korea.")
    p2 = doc.add_paragraph(style="MDPI_1.6_affiliation")
    p2.add_run("*").font.superscript = True
    p2.add_run("\tCorrespondence: [email-to-insert]")

    # Abstract
    p = doc.add_paragraph(style="MDPI_1.7_abstract")
    p.add_run("Abstract: ").bold = True
    p.add_run("Chronic insomnia affects more than 850 million adults worldwide "
              "and remains poorly served by long-term pharmacotherapy, motivating "
              "computational re-analysis of established traditional Chinese medicine "
              "(TCM) prescribing practice. Prescription corpora are high-dimensional "
              "and lexically noisy, so isolated frequency statistics rarely expose "
              "the underlying combinatorial logic, and a unified treatment that "
              "bridges local pair-level synergy with global network topology is "
              "still lacking. We assembled a de-identified, two-field transactional "
              "corpus of 14,178 (prescription, herb) records covering 986 distinct "
              "insomnia prescriptions and 326 raw herb expressions. A 24-rule "
              "ontology reduced 642 expressions (4.53%) to canonical pharmacopoeial "
              "names, and intra-prescription deduplication removed 28 redundant "
              "rows, yielding 986 prescriptions, 317 herb entities and 14,150 "
              "cleaned records (mean 14.35 ± 2.94 herbs per prescription; median 14; "
              "range 6–23). Frequent itemsets and binary association rules were "
              "extracted with the Apriori algorithm under the triple constraint "
              "min_support = 0.20, min_confidence = 0.60 and Lift > 1.0. A weighted, "
              "undirected co-occurrence graph (317 nodes, 11,239 edges) was then "
              "constructed, and three community-detection algorithms were applied "
              "in parallel. Thirty-four binary rules satisfied the triple constraint. "
              "The highest Lift values were observed for ")
    p.add_run("Bai Zhu").italic = True
    p.add_run(" – ")
    p.add_run("Fu Ling").italic = True
    p.add_run(" (Lift = 1.99) and ")
    p.add_run("Long Gu").italic = True
    p.add_run(" – ")
    p.add_run("Mu Li").italic = True
    p.add_run(" (1.84); the trio ")
    p.add_run("Long Gu – Mu Li – He Huan Pi").italic = True
    p.add_run(" formed a pairwise sub-graph in which every Lift exceeded 1.65. "
              "The co-occurrence network displayed small-world behaviour "
              "(mean clustering 0.83; diameter 3). The three algorithms each "
              "returned four functional modules with modularity Q ≈ 0.04, "
              "interpretable as Sedative–Tranquilising, Qi-Tonifying and "
              "Spleen-Strengthening, Heat-Clearing and Phlegm-Resolving, and "
              "Yin-Nourishing and Blood-Enriching clusters. All four flagship "
              "Lift ≥ 1.40 pairs persisted in 27/27 parameter cells of the "
              "sensitivity grid and in 100% of 200 Bootstrap resamples. The "
              "full dataset, ontology and analysis scripts are released as a "
              "single-command reproducible package; prospective randomised "
              "controlled trials are needed to translate the computational "
              "findings into clinical evidence.")

    # Keywords
    p = doc.add_paragraph(style="MDPI_1.8_keywords")
    p.add_run("Keywords: ").bold = True
    p.add_run("insomnia; traditional Chinese medicine; data mining; Apriori algorithm; "
              "association rules; complex network analysis; Louvain community detection; "
              "Leiden algorithm; reproducible research; computational pharmacology")

    # Page-rule separator
    _para(doc, "", "MDPI_1.9_line")

    # ── 1. Introduction ────────────────────────────────────────────
    _para(doc, "1. Introduction", "MDPI_2.1_heading1")
    _para(doc,
          "Insomnia is a chronic disturbance of sleep–wake regulation characterised "
          "by difficulty initiating or maintaining sleep, early-morning awakening, "
          "and unrefreshing rest. A recent systematic literature review estimated "
          "that more than 850 million adults globally meet a clinically meaningful "
          "insomnia definition, with a marked female and age-related skew [1]. "
          "The downstream burden—cognitive impairment, depression, cardiovascular "
          "risk, and elevated suicidality—has shifted insomnia from a quality-of-"
          "life issue to a frontline public-health concern.",
          "MDPI_3.1_text")
    _para(doc,
          "Cognitive-behavioural therapy for insomnia (CBT-I) is the first-line "
          "recommendation in both European and North American guidelines [2,3]. "
          "Long-term pharmacological control, however, remains contentious: a "
          "recent network meta-analysis questioned the durable efficacy of "
          "benzodiazepines and Z-drugs for chronic insomnia [4], and a parallel "
          "network synthesis of psychological and pharmacological options reached "
          "compatible conclusions [5]. Against this backdrop traditional Chinese "
          "medicine (TCM) has attracted growing computational attention. A global "
          "guideline mapping exercise documented the formal inclusion of TCM "
          "modalities for insomnia across multiple jurisdictions [6]; a network "
          "meta-analysis covering 109 randomised controlled trials reported tangible "
          "improvements in the Pittsburgh Sleep Quality Index and favourable safety "
          "profiles for several proprietary multi-herb formulations [7]; meta-"
          "analytic syntheses of Suanzaoren decoction [8], commercial Zao Ren An "
          "Shen preparations [32] and the single herb Suan Zao Ren [33] provide "
          "convergent evidence.",
          "MDPI_3.1_text")
    _para(doc,
          "Two methodological obstacles, however, have constrained the field. "
          "First, clinical TCM data are highly individualised: prescriptions "
          "follow physician-specific compositional logic, and the lexical surface "
          "is contaminated by synonyms, processing-method prefixes and provenance "
          "qualifiers, producing sparse high-dimensional records in which raw "
          "co-occurrence statistics frequently overstate the relevance of broadly "
          "used adjunct herbs [9,10]. Second, the literature is dominated by "
          "single-method designs: studies typically apply either association rule "
          "mining or clustering, rarely both, and rarely report the parameter "
          "sensitivity or sample-level stability that a reviewer would expect "
          "from a quantitative pipeline.",
          "MDPI_3.1_text")
    _para(doc,
          "The Apriori algorithm [11] remains the canonical engine for frequent-"
          "itemset discovery in sparse transactional data and has been widely "
          "deployed in disease-specific TCM mining studies [12]. Complementary "
          "insights come from graph-theoretic representations, in which the "
          "prescription corpus is encoded as a weighted graph and node centrality "
          "[39], modularity-optimising community detection [13,37] and Leiden "
          "refinement [14] expose global topology that itemset mining cannot. "
          "Recent work that combined network and clustering perspectives on "
          "insomnia symptom phenotypes [15] illustrates the value of multi-method "
          "designs.",
          "MDPI_3.1_text")
    _para(doc,
          "Building on this foundation, we report a dual-track pipeline that "
          "couples Apriori association mining with weighted-network analysis on "
          "a single de-identified insomnia prescription corpus. Five contributions "
          "distinguish the present work from earlier studies. First, the analysis "
          "rests on an ontology-normalised dataset of 986 prescriptions, 317 "
          "canonical herb entities and 14,150 cleaned records. Second, Apriori "
          "and weighted-network analysis are run in parallel and cross-validated. "
          "Third, parameter robustness is examined over a 3 × 3 × 3 grid (27 "
          "settings) and complemented by Bootstrap resampling. Fourth, the entire "
          "pipeline—eleven Python scripts, the conda environment and the "
          "manuscript source—is released as a single-command (\"make all\") "
          "reproducibility package. Fifth, every quantitative finding is mapped, "
          "with explicit caveats, onto both classical formulary theory and "
          "contemporary neuropharmacology.",
          "MDPI_3.1_text")

    # ── 2. Materials and Methods ───────────────────────────────────
    _para(doc, "2. Materials and Methods", "MDPI_2.1_heading1")

    _para(doc, "2.1. Data Sources", "MDPI_2.2_heading2")
    _para(doc,
          "Robust frequent-itemset extraction and reliable topology characterisation "
          "both require a high-quality, low-bias and geographically diverse corpus. "
          "We assembled an integrated, manually curated insomnia prescription "
          "dataset that draws on three classes of source without recourse to any "
          "institutional electronic medical record system.",
          "MDPI_3.1_text")
    _para(doc, "2.1.1. Source Categories", "MDPI_2.3_heading3")
    _para(doc,
          "Source provenance falls into three classes (an explicit inventory is "
          "provided in Supplementary Table S1). The first class comprises peer-"
          "reviewed clinical research and prescription literature on insomnia, "
          "retrieved from China National Knowledge Infrastructure (CNKI), Wanfang "
          "Data, VIP and PubMed; the second class consists of classical formulary "
          "databases—primarily the Database of Chinese Formulae maintained by the "
          "China National Population and Health Scientific Data Sharing Platform—"
          "together with historical canonical texts queried through period-specific "
          "synonyms for insomnia ('bu de wo', 'bu mei', 'mu bu ming'); the third "
          "class is a research-curated archive of clinical prescriptions that had "
          "already been de-identified before entering the present analysis, "
          "retaining only herb composition and a synthetic identifier.",
          "MDPI_3.1_text")
    _para(doc, "2.1.2. Field Minimisation and De-identification",
          "MDPI_2.3_heading3")
    p = doc.add_paragraph(style="MDPI_3.1_text")
    p.add_run("After cross-source consolidation, merging of duplicates and field "
              "standardisation, the dataset is stored in the long-format "
              "transactional file ")
    p.add_run("anonymized_prescription_transactions.csv").italic = True
    p.add_run(" containing only two atomic fields: ")
    p.add_run("prescription_id").italic = True
    p.add_run(" (an opaque, irreversibly mapped identifier) and ")
    p.add_run("herb").italic = True
    p.add_run(" (the original Chinese herb expression). No direct identifier "
              "(name, government identification number, contact details) and "
              "no quasi-identifier (visit date, institution-internal record "
              "number, physician identity, geographic indicator) is retained. "
              "The mapping table linking opaque identifiers to original sources "
              "is kept under encryption by the study team and is not distributed "
              "with the public dataset, so re-identification of any individual "
              "patient, prescriber or institution from the released file is not "
              "feasible by design.")
    _para(doc, "2.1.3. Inclusion and Exclusion Criteria",
          "MDPI_2.3_heading3")
    _para(doc,
          "Records were retained when the prescription explicitly targeted "
          "primary insomnia or the related TCM syndrome bu mei, the composition "
          "was complete and structurally extractable, and the source was "
          "traceable to a published article, a recognised database or a "
          "de-identified research archive. Records were excluded when the "
          "composition was incomplete or contained only a single herb, when "
          "the indication had no demonstrable link to insomnia, or when the "
          "same prescription appeared in more than one source (one instance "
          "retained).",
          "MDPI_3.1_text")
    _para(doc, "2.1.4. Final Dataset Size", "MDPI_2.3_heading3")
    p = doc.add_paragraph(style="MDPI_3.1_text")
    p.add_run("The unfiltered transactional table comprised ")
    p.add_run("14,178").bold = True
    p.add_run(" (prescription, herb) records spanning ")
    p.add_run("986").bold = True
    p.add_run(" prescriptions and ")
    p.add_run("326").bold = True
    p.add_run(" raw herb expressions; the ")
    p.add_run("prescription_id").italic = True
    p.add_run(" and ")
    p.add_run("herb").italic = True
    p.add_run(" fields had no missing values and no fully duplicated rows. After "
              "the preprocessing steps described in Section 2.2—ontology-driven "
              "normalisation (24 rules; 642 records merged, equivalent to ")
    p.add_run("4.53%").bold = True
    p.add_run(" of the raw table) and intra-prescription deduplication (28 "
              "redundant rows removed)—the dataset that entered Apriori and "
              "network analysis consisted of ")
    p.add_run("986 prescriptions, 317 canonical herb entities and 14,150 "
              "transactional records").bold = True
    p.add_run(", with a mean of ")
    p.add_run("14.35 ± 2.94").bold = True
    p.add_run(" herbs per prescription (median 14; range 6–23 herbs).")

    _para(doc, "2.2. Preprocessing and Ontology Normalisation",
          "MDPI_2.2_heading2")
    p = doc.add_paragraph(style="MDPI_3.1_text")
    p.add_run("Preprocessing proceeded in four steps. (i) Completeness audit: "
              "both fields were screened for missing values and for fully "
              "duplicated ")
    p.add_run("(prescription_id, herb)").italic = True
    p.add_run(" pairs; the count of each was zero. (ii) Ontology-driven "
              "normalisation: 24 mapping rules consolidating processing-method "
              "suffixes (e.g., ")
    p.add_run("chao bai zhu → bai zhu").italic = True
    p.add_run("), pharmacopoeial synonyms (")
    p.add_run("ye jiao teng → shou wu teng").italic = True
    p.add_run("), provenance prefixes (")
    p.add_run("chuan niu xi / huai niu xi → niu xi").italic = True
    p.add_run(") and preparation prefixes (")
    p.add_run("cu xiang fu → xiang fu").italic = True
    p.add_run(") were applied. The mapping consolidated 642 records (4.5281%) "
              "and reduced the herb-expression vocabulary from 326 to 317. "
              "(iii) Intra-prescription deduplication: 28 rows in which a "
              "normalised herb appeared more than once within the same "
              "prescription were collapsed. (iv) Quality control: a second pass "
              "confirmed that no residual unnormalised aliases remained. Pre- "
              "and post-cleaning counts are summarised in Table 1.")

    # Table 1
    _para(doc, "Table 1. Dataset metrics before and after preprocessing.",
          "MDPI_4.1_table_caption")
    _table(doc, [
        ["Metric", "Pre-cleaning", "Post-cleaning", "Note"],
        ["Prescriptions", "986", "986", "Unchanged"],
        ["Herb expressions", "326", "317",
         "Nine consolidated by the 24-rule ontology"],
        ["Transactional records", "14,178", "14,150",
         "28 intra-prescription redundancies removed"],
        ["Records affected by normalisation", "—", "642",
         "4.53% of the raw table"],
        ["Residual unnormalised aliases", "—", "0",
         "Second-pass quality check"],
        ["Herbs per prescription (mean ± SD)", "14.38 ± 2.94",
         "14.35 ± 2.94", "Median 14; range 6–23"],
        ["Fully duplicated (id, herb) rows", "0", "0", "None in either state"],
    ], col_widths_cm=[5.5, 3.0, 3.2, 6.8])

    _para(doc, "2.3. Prescription–Herb Transaction Matrix",
          "MDPI_2.2_heading2")
    p = doc.add_paragraph(style="MDPI_3.1_text")
    p.add_run("We encoded the cleaned corpus as a Boolean transaction matrix ")
    p.add_run("M ∈ {0, 1}").italic = True
    p.add_run("986 × 317").font.superscript = True
    p.add_run(", with row ")
    p.add_run("i").italic = True
    p.add_run(" denoting a prescription, column ")
    p.add_run("j").italic = True
    p.add_run(" a canonical herb and ")
    p.add_run("M[i, j] = 1").italic = True
    p.add_run(" when herb ")
    p.add_run("j").italic = True
    p.add_run(" was present in prescription ")
    p.add_run("i").italic = True
    p.add_run(". The column-wise sum gives the global frequency of each herb. "
              "The empirical herb frequency distribution exhibits a long tail: "
              "a Pareto-style cumulative-coverage analysis (Figure 2) showed "
              "that 43 herbs (13.56% of all canonical entities) together account "
              "for 80% of the cumulative usage, providing a natural pruning "
              "horizon for both Apriori and downstream network analysis.")

    _para(doc, "2.4. Apriori Association Rule Mining",
          "MDPI_2.2_heading2")
    p = doc.add_paragraph(style="MDPI_3.1_text")
    p.add_run(
        "Frequent itemsets up to size four were generated with the Apriori "
        "algorithm [11]. We evaluated rules on three standard metrics: support, "
        "Support(X → Y) = |X ∪ Y| / N; confidence, Confidence(X → Y) = "
        "Support(X ∪ Y) / Support(X); and Lift, Lift(X → Y) = "
        "Confidence(X → Y) / Support(Y). A Lift exceeding unity indicates a "
        "positive co-occurrence beyond chance, whereas Lift = 1 corresponds "
        "to statistical independence; including Lift in the filter is "
        "essential here because indiscriminate harmoniser herbs (notably "
    )
    p.add_run("Gan Cao").italic = True
    p.add_run(") generate inflated confidence values that are not "
              "pharmacologically informative.")
    _para(doc,
          "Baseline thresholds were min_support = 0.20 (the joint occurrence "
          "had to involve at least ≈ 197 prescriptions), min_confidence = 0.60 "
          "and Lift > 1.0. The algorithm was implemented through the apriori "
          "and association_rules functions of mlxtend 0.23.0 [40] running on "
          "Python 3.11. To assess the influence of these thresholds we re-ran "
          "the procedure on a full 3 × 3 × 3 grid—min_support ∈ {0.15, 0.20, "
          "0.25}, min_confidence ∈ {0.50, 0.60, 0.70} and Lift threshold ∈ "
          "{1.0, 1.2, 1.4}—and recorded both the number of surviving binary "
          "rules and the persistence of four flagship strong-Lift pairs "
          "(Section 3.7 and Table 6).",
          "MDPI_3.1_text")

    _para(doc, "2.5. Co-occurrence Network and Centrality",
          "MDPI_2.2_heading2")
    p = doc.add_paragraph(style="MDPI_3.1_text")
    p.add_run("We represented the prescription corpus as a weighted, undirected "
              "graph ")
    p.add_run("G = (V, E)").italic = True
    p.add_run(". Each node ")
    p.add_run("v ∈ V").italic = True
    p.add_run(" is one of the 317 canonical herbs; an edge connects two herbs "
              "whenever they co-occur in at least one prescription, and the edge "
              "weight equals the number of prescriptions in which they co-appear. "
              "Each edge additionally carries a Lift attribute, permitting "
              "downstream Lift-weighted filtering. We computed degree, "
              "betweenness and closeness centrality [39] with NetworkX 3.1 [36], "
              "and assessed small-world organisation [38] through the mean "
              "clustering coefficient, the mean shortest-path length and the "
              "diameter, contrasting the observed values with size- and "
              "density-matched Erdős–Rényi controls.")

    _para(doc, "2.6. Community Detection", "MDPI_2.2_heading2")
    p = doc.add_paragraph(style="MDPI_3.1_text")
    p.add_run("Three modularity-optimising community-detection algorithms were "
              "applied in parallel. The Louvain method [13] served as the primary "
              "algorithm; the Greedy Modularity heuristic implemented in NetworkX "
              "[37] and the Leiden algorithm [14] (CPM quality function, "
              "resolution parameter γ = 1.0, ")
    p.add_run("leidenalg").italic = True
    p.add_run(" 0.11) served as comparators. Random seeds were fixed at 42 "
              "throughout. Modularity ")
    p.add_run("Q").italic = True
    p.add_run(" is reported per algorithm, and pairwise agreement is quantified "
              "through Adjusted Rand Index (ARI) and Normalised Mutual Information "
              "(NMI) using ")
    p.add_run("scikit-learn").italic = True
    p.add_run(" 1.3.0.")

    _para(doc, "2.7. Reproducibility and Computing Environment",
          "MDPI_2.2_heading2")
    _para(doc,
          "All computations were performed in Python 3.11.6 on Ubuntu 22.04 "
          "(Intel Xeon Gold, 64 GB RAM). Core dependencies were pandas 2.0.3, "
          "numpy 1.24.4, scipy 1.11.2, mlxtend 0.23.0 [40], networkx 3.1 [36], "
          "python-louvain 0.16, leidenalg 0.11, scikit-learn 1.3.0, matplotlib "
          "3.7.2 and seaborn 0.12.2. Eleven ordered scripts (step1_data_audit.py "
          "through step11_reproducibility_report.py) drive the pipeline end-to-"
          "end, write thirteen intermediate CSV / JSON artefacts, six tables and "
          "eight figures, and generate an HTML reproducibility report that "
          "asserts every headline number against the manuscript values. The "
          "package is released under MIT (code) and CC-BY 4.0 (data, tables, "
          "figures), and is permanently archived on Zenodo with the DOI cited "
          "in the Data Availability Statement.",
          "MDPI_3.1_text")

    # ── 3. Results ─────────────────────────────────────────────────
    _para(doc, "3. Results", "MDPI_2.1_heading1")

    _para(doc, "3.1. Descriptive Statistics", "MDPI_2.2_heading2")
    _para(doc,
          "The cleaned dataset comprises 986 prescriptions and 317 canonical "
          "herb entities, totalling 14,150 transactional records. Prescription "
          "length follows a near-normal distribution (mean 14.35 ± 2.94; median "
          "14; range 6–23), consistent with reported norms for insomnia "
          "prescribing in routine TCM practice.",
          "MDPI_3.1_text")

    _para(doc, "3.2. Herb Frequency", "MDPI_2.2_heading2")
    _para(doc,
          "Frequency was dominated by the two compounds that anchor classical "
          "insomnia formularies (Table 2). Gan Cao (Glycyrrhizae Radix et "
          "Rhizoma) and Suan Zao Ren (Ziziphi Spinosae Semen) together accounted "
          "for 11.11% of the cumulative usage; six of the top ten herbs were "
          "nervine sedatives, including Long Gu (Os Draconis) and Mu Li "
          "(Ostreae Concha) as mineral-class sedatives. Nine of the top ten "
          "herbs are reported to enter the Heart meridian and five enter the "
          "Liver meridian, a distribution that closely mirrors the classical "
          "rationale that the Heart houses the Spirit.",
          "MDPI_3.1_text")
    _para(doc, "Table 2. Ten most frequent canonical herbs.",
          "MDPI_4.1_table_caption")
    _table(doc, [
        ["Rank", "Herb (Pinyin)", "Latin pharmacognostic name",
         "Frequency", "Prevalence (%)", "Cumulative (%)"],
        ["1",  "Gan Cao",      "Glycyrrhizae Radix et Rhizoma", "806", "81.74", "5.70"],
        ["2",  "Suan Zao Ren", "Ziziphi Spinosae Semen",        "766", "77.69", "11.11"],
        ["3",  "He Huan Pi",   "Albiziae Cortex",               "503", "51.01", "14.66"],
        ["4",  "Shou Wu Teng", "Polygoni Multiflori Caulis",    "501", "50.81", "18.20"],
        ["5",  "Fu Shen",      "Poria cum Radice Pini",         "462", "46.86", "21.47"],
        ["6",  "Yuan Zhi",     "Polygalae Radix",               "461", "46.75", "24.73"],
        ["7",  "Wu Wei Zi",    "Schisandrae Chinensis Fructus", "437", "44.32", "27.82"],
        ["8",  "Long Gu",      "Os Draconis",                   "429", "43.51", "30.85"],
        ["9",  "Mu Li",        "Ostreae Concha",                "425", "43.10", "33.85"],
        ["10", "Fu Ling",      "Poria",                         "422", "42.80", "36.83"],
    ], col_widths_cm=[1.3, 2.5, 5.5, 2.0, 2.4, 2.8])

    _figure(doc, FIG / "figure2_pareto_top30.png",
            "Figure 2. Pareto distribution of the 30 most frequent herbs. "
            "The orange bars highlight the ten leading herbs; the red line "
            "marks the cumulative share of total usage, with the dashed reference "
            "at 80%. Together the top thirty herbs account for 67.74% of the "
            "cumulative usage.")

    _para(doc, "3.3. Apriori Association Rules", "MDPI_2.2_heading2")
    _para(doc,
          "At the baseline triple constraint the algorithm produced 136 "
          "frequent itemsets (19 one-, 52 two-, 51 three- and 14 four-itemsets). "
          "After confidence and Lift filtering, 247 rules survived, 34 of which "
          "were binary. Three behavioural archetypes emerge from the binary set. "
          "The high-support–low-Lift archetype is exemplified by Suan Zao Ren ↔ "
          "Gan Cao (support 64.30%, Lift 1.01), where the joint occurrence "
          "reflects ubiquitous use of both rather than pharmacological pairing. "
          "The high-confidence directional archetype is exemplified by Wu Wei "
          "Zi → Suan Zao Ren (confidence 90.16%) and Fu Shen → Suan Zao Ren "
          "(90.04%), in which Suan Zao Ren operates as the central sedative "
          "hub. The high-Lift functional archetype is exemplified by Bai Zhu ↔ "
          "Fu Ling (Lift 1.99) and Long Gu ↔ Mu Li (Lift 1.84), in which the "
          "two herbs preferentially co-appear well above the chance expectation. "
          "Table 3 reports a representative selection from the 34 binary rules.",
          "MDPI_3.1_text")
    _para(doc,
          "Table 3. Selected Apriori binary association rules (full table in "
          "Supplementary Materials).",
          "MDPI_4.1_table_caption")
    _table(doc, [
        ["LHS", "RHS", "Support (%)", "Confidence (%)", "Lift"],
        ["Gan Cao",      "Suan Zao Ren", "64.30", "78.66", "1.01"],
        ["Suan Zao Ren", "Gan Cao",      "64.30", "82.77", "1.01"],
        ["Shou Wu Teng", "Suan Zao Ren", "44.52", "87.62", "1.13"],
        ["He Huan Pi",   "Suan Zao Ren", "42.70", "83.70", "1.08"],
        ["Fu Shen",      "Suan Zao Ren", "42.19", "90.04", "1.16"],
        ["Yuan Zhi",     "Suan Zao Ren", "41.89", "89.59", "1.15"],
        ["Wu Wei Zi",    "Suan Zao Ren", "39.96", "90.16", "1.16"],
        ["Shou Wu Teng", "He Huan Pi",   "38.64", "76.05", "1.49"],
        ["Mu Li",        "He Huan Pi",   "37.12", "86.12", "1.69"],
        ["Long Gu",      "He Huan Pi",   "36.71", "84.38", "1.65"],
        ["Long Gu",      "Mu Li",        "34.58", "79.49", "1.84"],
        ["Bai Zhu",      "Fu Ling",      "29.31", "85.00", "1.99"],
        ["Bai Zi Ren",   "Yuan Zhi",     "25.66", "72.70", "1.55"],
    ], col_widths_cm=[3.4, 3.4, 3.0, 3.0, 3.7])
    _figure(doc, FIG / "figure3_apriori_bubble.png",
            "Figure 3. Apriori binary association rules in the support–"
            "confidence–lift space. Marker size scales with Lift; rules with "
            "Lift ≥ 1.40 are annotated with the herb pair and the Lift value.")

    _para(doc, "3.4. Network Topology", "MDPI_2.2_heading2")
    _para(doc,
          "The full co-occurrence graph carried 317 nodes and 11,239 edges, "
          "with a mean degree of 70.91, a mean clustering coefficient of 0.83, "
          "a network diameter of 3 and a density of 0.224. The clustering–path-"
          "length profile is consistent with classical small-world organisation "
          "[38]. Figure 4 visualises the Top-30 sub-graph after Lift filtering "
          "at the 1.05 threshold, which suppresses non-specific connections "
          "involving universal harmoniser herbs and exposes two tightly knit "
          "modules. One module is built around Suan Zao Ren, He Huan Pi, Shou "
          "Wu Teng, Fu Shen, Yuan Zhi, Wu Wei Zi, Long Gu and Mu Li; the other "
          "groups Gan Cao, Fu Ling, Bai Zhu, Dang Gui, Dang Shen, Ren Shen, "
          "Long Yan Rou, Mu Xiang, Da Zao and Huang Qi.",
          "MDPI_3.1_text")
    _figure(doc, FIG / "figure4_core_network.png",
            "Figure 4. Core co-occurrence network for the 30 most frequent "
            "herbs. Edges with Lift below 1.05 have been suppressed to "
            "emphasise structurally informative connections. Node size scales "
            "with raw frequency.")

    _para(doc, "3.5. Centrality Profiles", "MDPI_2.2_heading2")
    _para(doc,
          "Degree and closeness centrality identified Gan Cao (DC = 0.987; CC "
          "= 0.988) and Suan Zao Ren (DC = 0.968; CC = 0.969) as the network's "
          "twin hubs, followed by He Huan Pi, Shou Wu Teng, Wu Wei Zi, Fu Shen, "
          "Mu Li, Yuan Zhi, Fu Ling and Long Gu (Table 4). Betweenness "
          "centrality, in contrast, returned a more dispersed Top-10 led by "
          "Long Yan Rou (BC = 0.0192), Mu Xiang (0.0181) and Shi Chang Pu "
          "(0.0166)—moderate-frequency herbs that nevertheless lie on the "
          "shortest paths bridging functional modules. We refer to these as "
          "candidate bridge nodes (Figure 5). A pairwise Lift heatmap restricted "
          "to the top fifteen herbs (Figure 6) confirms that the high-Lift "
          "signal is concentrated in six pairs: Bai Zhu – Fu Ling (1.99), "
          "Long Gu – Mu Li (1.84), He Huan Pi – Mu Li (1.69), He Huan Pi – "
          "Long Gu (1.65), Bai Zi Ren – Yuan Zhi (1.55) and He Huan Pi – "
          "Shou Wu Teng (1.49).",
          "MDPI_3.1_text")
    _para(doc, "Table 4. Top ten herbs by degree (DC), betweenness (BC) and "
                "closeness (CC) centrality.", "MDPI_4.1_table_caption")
    _table(doc, [
        ["Rank", "DC herb", "DC", "BC herb", "BC", "CC herb", "CC"],
        ["1",  "Gan Cao",      "0.987", "Long Yan Rou", "0.0192", "Gan Cao",      "0.988"],
        ["2",  "Suan Zao Ren", "0.968", "Mu Xiang",     "0.0181", "Suan Zao Ren", "0.969"],
        ["3",  "He Huan Pi",   "0.911", "Shi Chang Pu", "0.0166", "He Huan Pi",   "0.919"],
        ["4",  "Shou Wu Teng", "0.908", "Nv Zhen Zi",   "0.0166", "Shou Wu Teng", "0.916"],
        ["5",  "Wu Wei Zi",    "0.902", "Bai Shao",     "0.0163", "Wu Wei Zi",    "0.911"],
        ["6",  "Fu Shen",      "0.899", "Dang Shen",    "0.0161", "Fu Shen",      "0.908"],
        ["7",  "Mu Li",        "0.886", "Bei Sha Shen", "0.0146", "Mu Li",        "0.898"],
        ["8",  "Yuan Zhi",     "0.877", "Huang Qi",     "0.0146", "Yuan Zhi",     "0.890"],
        ["9",  "Fu Ling",      "0.867", "Ban Xia",      "0.0145", "Fu Ling",      "0.883"],
        ["10", "Long Gu",      "0.858", "Chen Pi",      "0.0144", "Long Gu",      "0.875"],
    ], col_widths_cm=[1.2, 2.6, 1.7, 2.6, 2.0, 2.6, 1.8])
    _figure(doc, FIG / "figure5_centrality_top10.png",
            "Figure 5. Degree, betweenness and closeness centrality for the "
            "ten most frequent herbs. Betweenness values have been multiplied "
            "by 30 to allow visual comparison with the other two normalised "
            "centralities.")
    _figure(doc, FIG / "figure6_lift_heatmap.png",
            "Figure 6. Pairwise Lift heatmap for the fifteen most frequent "
            "herbs. Cells with Lift ≥ 1.40 are outlined in black: Bai Zhu – "
            "Fu Ling (1.99); Long Gu – Mu Li (1.84); He Huan Pi – Mu Li "
            "(1.69); He Huan Pi – Long Gu (1.65); Bai Zi Ren – Yuan Zhi "
            "(1.55); He Huan Pi – Shou Wu Teng (1.49).")

    _para(doc, "3.6. Community Structure", "MDPI_2.2_heading2")
    _para(doc,
          "All three algorithms (Louvain, Greedy Modularity, Leiden) returned "
          "four communities on the full 317-node graph, with respective "
          "modularities of 0.0378, 0.0361 and 0.0385. The relatively low Q "
          "values reflect the pervasive co-occurrence of universal harmonisers "
          "(Gan Cao, Suan Zao Ren) with virtually every functional class, "
          "which softens the inter-community boundary in the raw co-occurrence "
          "space. Despite this numerical compression, the four modules carry "
          "coherent functional interpretations (Table 5) that align with the "
          "classical syndrome categorisation of insomnia: a Sedative–"
          "Tranquilising cluster, a Qi-Tonifying and Spleen-Strengthening "
          "cluster, a Heat-Clearing and Phlegm-Resolving cluster, and a Yin-"
          "Nourishing and Blood-Enriching cluster.",
          "MDPI_3.1_text")
    _para(doc,
          "Table 5. Four functional communities recovered by Louvain on the "
          "full graph.", "MDPI_4.1_table_caption")
    _table(doc, [
        ["Module", "Core members (descending frequency)", "Interpretation",
         "Plausible syndrome", "Data support"],
        ["A. Sedative–Tranquilising",
         "Suan Zao Ren, He Huan Pi, Shou Wu Teng, Fu Shen, Yuan Zhi, "
         "Wu Wei Zi, Long Gu, Mu Li, Bai Zi Ren, Shi Chang Pu",
         "Nourishment of the heart and tranquilisation of the spirit",
         "Heart-spirit disquiet; liver–heart blood deficiency",
         "Strong (He Huan Pi – Mu Li, Long Gu – Mu Li, Bai Zi Ren – "
         "Yuan Zhi carry Lift ≥ 1.40)"],
        ["B. Qi-Tonifying and Spleen-Strengthening",
         "Gan Cao, Fu Ling, Dang Gui, Bai Zhu, Long Yan Rou, Huang Qi, "
         "Dang Shen, Da Zao, Mu Xiang, Ren Shen",
         "Replenishment of qi and fortification of the spleen",
         "Concurrent heart and spleen deficiency",
         "Strong (Bai Zhu – Fu Ling Lift = 1.99)"],
        ["C. Heat-Clearing and Phlegm-Resolving",
         "Huang Bai, Ban Xia, Zhu Ru, Mu Dan Pi, Zhi Shi, Huang Lian, Zhi Zi, "
         "Hou Po, Huang Qin, Chai Hu",
         "Clearance of heat with phlegm resolution",
         "Phlegm-heat disturbance; liver constraint progressing to fire",
         "Partial (no within-module Lift ≥ 1.40 pair in the Top-15 frame)"],
        ["D. Yin-Nourishing and Blood-Enriching",
         "Bai Shao, Sheng Di Huang, Mai Dong, Bei Sha Shen, Chuan Xiong, "
         "Zhi Mu, Nv Zhen Zi, Yu Zhu, Shan Zhu Yu, E Jiao",
         "Yin replenishment with heat clearance",
         "Yin deficiency with effulgent fire; liver–kidney depletion",
         "Partial (pairwise Lifts cluster near unity)"],
    ], col_widths_cm=[3.7, 5.0, 3.5, 3.5, 2.5])

    _para(doc, "3.7. Parameter Sensitivity", "MDPI_2.2_heading2")
    _para(doc,
          "We exhaustively swept the 3 × 3 × 3 grid of (min_support, "
          "min_confidence, Lift threshold). At every one of the 27 settings "
          "the four flagship pairs—Bai Zhu – Fu Ling, Long Gu – Mu Li, "
          "He Huan Pi – Mu Li and He Huan Pi – Long Gu—survived; the Top-10 "
          "frequency ordering was identical to the baseline ordering across "
          "every cell of the grid. A representative subset of the grid is "
          "presented in Table 6 and the full surface is shown in Figure 7.",
          "MDPI_3.1_text")
    _para(doc,
          "Table 6. Apriori sensitivity over a representative 9-point subset "
          "of the 27-cell grid.", "MDPI_4.1_table_caption")
    _table(doc, [
        ["min_support", "min_confidence", "Lift threshold", "Binary rules",
         "Flagship pairs", "Top-10 ordering"],
        ["0.15", "0.50", "1.0", "79", "4 / 4", "concordant"],
        ["0.15", "0.60", "1.0", "42", "4 / 4", "concordant"],
        ["0.15", "0.70", "1.0", "39", "4 / 4", "concordant"],
        ["0.20", "0.50", "1.0", "49", "4 / 4", "concordant"],
        ["0.20", "0.60", "1.0 (baseline)", "34", "4 / 4", "concordant"],
        ["0.20", "0.60", "1.4", "9",  "4 / 4", "concordant"],
        ["0.25", "0.50", "1.0", "28", "4 / 4", "concordant"],
        ["0.25", "0.60", "1.0", "19", "4 / 4", "concordant"],
        ["0.25", "0.60", "1.4", "7",  "4 / 4", "concordant"],
    ], col_widths_cm=[2.6, 3.0, 2.7, 2.5, 2.5, 4.5])
    _figure(doc, FIG / "figure7_sensitivity.png",
            "Figure 7. Number of binary rules across the 27-cell sensitivity "
            "grid. Lines distinguish confidence thresholds, marker size "
            "encodes the Lift threshold and the horizontal axis represents "
            "min_support.")

    _para(doc, "3.8. Bootstrap Stability and Algorithmic Concordance",
          "MDPI_2.2_heading2")
    _para(doc,
          "Two hundred Bootstrap resamples drawn at 80% prescription coverage "
          "were used to probe sample-level stability. Top-10 herb membership "
          "matched the baseline in 100% of the resamples; each of the four "
          "flagship Lift ≥ 1.40 pairs persisted in 100% of the resamples, as "
          "did the two secondary pairs Bai Zi Ren – Yuan Zhi and He Huan Pi – "
          "Shou Wu Teng. The three community-detection algorithms reached "
          "pairwise ARI values of 0.41 (Louvain–Greedy), 0.51 (Louvain–Leiden) "
          "and 0.46 (Greedy–Leiden), with parallel NMI values of 0.40, 0.48 "
          "and 0.44 (Figure 8). The moderate ARI / NMI levels reflect the "
          "soft module boundaries already noted in Section 3.6; the algorithms "
          "agree on the four-module partition at the global scale but place "
          "individual peripheral herbs differently.",
          "MDPI_3.1_text")
    _figure(doc, FIG / "figure8_community_consistency.png",
            "Figure 8. Pairwise Adjusted Rand Index and Normalised Mutual "
            "Information across the three community-detection algorithms.")

    # ── 4. Discussion ──────────────────────────────────────────────
    _para(doc, "4. Discussion", "MDPI_2.1_heading1")
    _para(doc, "4.1. Frequency Profile and the Classical Heart-Spirit Rationale",
          "MDPI_2.2_heading2")
    _para(doc,
          "The two herbs that top the frequency table—Gan Cao and Suan Zao "
          "Ren—are formulary archetypes: Gan Cao is the canonical harmoniser "
          "introduced into almost every multi-herb prescription, whereas "
          "Suan Zao Ren is the sovereign ingredient of Suan Zao Ren decoction "
          "(Jin Gui Yao Lue, Han dynasty). The remaining sedative members of "
          "the Top-10—He Huan Pi, Shou Wu Teng, Fu Shen, Yuan Zhi, Long Gu "
          "and Mu Li—collectively account for six of the ten leading entries "
          "and define the operational core of insomnia therapy in this corpus. "
          "The dominance of Heart-meridian herbs (9 / 10) is consistent with "
          "the classical aphorism that the Heart houses the Spirit. Module B "
          "reproduces the canonical composition of Gui Pi Tang (Yan Yong-He, "
          "Ji Sheng Fang, Song dynasty), and the systematic synthesis of 109 "
          "RCTs reviewed in [7] reports that Gui Pi Tang-type formulations "
          "confer measurable benefit in patients with concurrent heart–spleen "
          "deficiency.",
          "MDPI_3.1_text")

    _para(doc, "4.2. Three Archetypes of Pair-Level Synergy",
          "MDPI_2.2_heading2")
    _para(doc,
          "Within the 34 binary rules we recovered three behavioural archetypes. "
          "The first, exemplified by Suan Zao Ren ↔ Gan Cao (support 64.30%, "
          "Lift 1.01), is a universal-pairing archetype: the two herbs are "
          "ubiquitous, so their joint presence in two-thirds of all prescriptions "
          "does not represent pharmacological synergy. The second, exemplified "
          "by Bai Zhu ↔ Fu Ling (Lift 1.99), is a high-Lift functional-pair "
          "archetype, in which two herbs preferentially co-occur. Bai Zhu "
          "fortifies the spleen and dries dampness; Fu Ling drains dampness "
          "and quietens the spirit; jointly the pair forms the dampness-clearing "
          "core of Gui Pi Tang and supplies the qi–blood substrate that anchors "
          "the broader formulation. The third archetype is the three-herb "
          "mineral-tranquilising sub-graph composed of Long Gu, Mu Li and He "
          "Huan Pi, in which every pairwise Lift exceeds 1.65. Long Gu and "
          "Mu Li are the two principal mineral sedatives in classical "
          "formularies; He Huan Pi adds a vegetal anxiolytic colouring. We "
          "additionally recover Bai Zi Ren ↔ Yuan Zhi (Lift 1.55), a balanced "
          "moistening–opening pair often invoked under the classical heuristic "
          "of complementary natures with reinforcing functions.",
          "MDPI_3.1_text")

    _para(doc, "4.3. Cross-Scale Mapping to Contemporary Neuropharmacology",
          "MDPI_2.2_heading2")
    _para(doc,
          "Modern mechanistic studies offer partial molecular correlates for "
          "the modules we recovered, though they cannot establish causal "
          "correspondence with the statistical pairings. For the Sedative–"
          "Tranquilising module, integrated network-pharmacology and molecular-"
          "docking analyses of Suan Zao Ren decoction [18] and UPLC-Q-TOF-MS "
          "multi-tissue metabolomics in p-chlorophenylalanine insomnia rats "
          "[19] suggest multi-pathway, multi-target sedative effects centred "
          "on Suan Zao Ren. Mechanistic work has further implicated flavonoid "
          "restructuring of the gut–brain axis in chronic restraint stress "
          "[20], GABA-A receptor activation [21] and modulation of the "
          "hypothalamic GABA / Glu balance with concomitant changes in "
          "Lactobacillus johnsonii abundance [35]; single-component studies of "
          "Ziziphi Spinosae oil terpenoids [22] and the saponin jujuboside A "
          "acting on the paraventricular thalamic GABAergic system [34] are "
          "mutually compatible. For Yuan Zhi, the protective and HPA-axis "
          "modulating activities reviewed in [23] and the multi-neurotransmitter "
          "shifts in the Suan Zao Ren–Yuan Zhi pair [24] provide an indirect "
          "mechanistic context for the high-confidence Suan Zao Ren-centred "
          "edges that dominate Section 3.3. A recent biophysical study of "
          "Long Gu-derived nanoparticles [25] proposes a calcium-dependent "
          "enterochromaffin / vagal-NTS pathway underpinning the mineral-"
          "tranquilising trio of Long Gu, Mu Li and He Huan Pi. For the "
          "He Huan Pi – Shou Wu Teng pair (Lift 1.49), immobilised 5-HT1A "
          "receptor affinity chromatography [26] identifies catechin and "
          "2,3,5,4'-tetrahydroxystilbene-2-O-β-D-glucoside as candidate active "
          "ligands. The GABAergic axis itself remains a parsimonious common "
          "substrate at the systems level [27], and integrative analyses "
          "combining network pharmacology with public expression repositories "
          "illustrate the multi-target nature of TCM insomnia treatments in "
          "general [28]. We stress that these molecular correlates remain "
          "plausibility arguments: the present study uses no biological assays "
          "and reports no causal inference.",
          "MDPI_3.1_text")

    _para(doc, "4.4. Methodological Position and Future Directions",
          "MDPI_2.2_heading2")
    _para(doc,
          "Three methodological choices distinguish this work from prior "
          "insomnia-TCM mining studies. We coupled Apriori with weighted-network "
          "analysis rather than relying on either family alone; we tested the "
          "durability of the headline findings under a 27-point parameter grid "
          "and 200 Bootstrap resamples; and we packaged the entire pipeline as "
          "a one-command, version-controlled reproducibility artefact. The "
          "pipeline can be extended in several directions. Dose-weighted "
          "association mining—which our Boolean transactional encoding does "
          "not support—could resolve the question of whether functional pair-"
          "level synergies are dose-modulated [16]. Multi-graph convolutional "
          "networks [30] and graph-neural prescription similarity [31] can "
          "carry the present static topology into a learned prescription-"
          "recommendation framework, and recent multi-graph embedding pipelines "
          "with data augmentation [17] show how data sparsity can be mitigated "
          "in this direction. The AI-assisted literature exploration approaches "
          "summarised in [29] offer a route to scaling the de-identification "
          "ontology and the bibliometric backbone used in the present study [10].",
          "MDPI_3.1_text")

    _para(doc, "4.5. Limitations", "MDPI_2.2_heading2")
    _para(doc,
          "Several caveats apply. (i) The Boolean transactional encoding "
          "ignores dose: two herbs paired at high frequency but radically "
          "different dosages collapse to identical co-occurrences. (ii) Patient-"
          "level syndrome differentiation and four-pillar diagnostic information "
          "are not represented in the corpus, limiting prescription–phenotype "
          "mapping. (iii) The low absolute modularity Q (≈ 0.04) reflects the "
          "genuinely soft community boundaries imposed by universal harmoniser "
          "herbs; future work could employ Lift-weighted or k-core refinement "
          "to sharpen these boundaries [14]. (iv) The strong synergistic "
          "combinations identified here (Bai Zhu – Fu Ling, Long Gu – Mu Li – "
          "He Huan Pi, Bai Zi Ren – Yuan Zhi) are computational candidates "
          "rather than validated therapeutic recommendations: prospective "
          "randomised controlled trials are needed to establish clinical "
          "efficacy and safety, in line with the cautionary stance taken by "
          "recent insomnia-TCM systematic reviews [7,8]. (v) The corpus "
          "represents a de-identified curated sample rather than an unbiased "
          "clinical census; external validation on larger, multi-centre samples "
          "is desirable before generalising the findings.",
          "MDPI_3.1_text")

    # ── 5. Conclusions ─────────────────────────────────────────────
    _para(doc, "5. Conclusions", "MDPI_2.1_heading1")
    _para(doc,
          "A dual-track pipeline that pairs Apriori association rule mining "
          "with weighted-network analysis was applied to 986 de-identified "
          "insomnia prescriptions covering 317 canonical herbs and 14,150 "
          "transactional records. Frequency profiling and centrality analysis "
          "confirmed a sedative-dominated prescribing pattern anchored by the "
          "twin hubs Gan Cao and Suan Zao Ren. Thirty-four binary association "
          "rules were recovered under the triple constraint, with Bai Zhu – "
          "Fu Ling (Lift = 1.99) and Long Gu – Mu Li (Lift = 1.84) ranking "
          "highest by Lift and Wu Wei Zi → Suan Zao Ren highest by directional "
          "confidence (90.16%). A 27-cell sensitivity grid and 200-iteration "
          "Bootstrap resampling demonstrated that the four flagship Lift ≥ 1.40 "
          "pairs and the Top-10 frequency ordering are highly stable across "
          "parameter and sample perturbation. The Louvain, Greedy Modularity "
          "and Leiden algorithms each returned four functional communities—"
          "Sedative–Tranquilising, Qi-Tonifying and Spleen-Strengthening, "
          "Heat-Clearing and Phlegm-Resolving, and Yin-Nourishing and Blood-"
          "Enriching—that map cleanly onto the classical syndrome "
          "differentiation of insomnia in traditional Chinese medicine. The "
          "high-Lift synergistic sub-structures (Bai Zhu – Fu Ling; Long Gu – "
          "Mu Li – He Huan Pi; Bai Zi Ren – Yuan Zhi) constitute quantitative "
          "restatements of the formulary heuristics of fortify-the-spleen and "
          "drain dampness, mineral-class tranquilisation and moistening–opening "
          "complementarity respectively, and can serve as candidate building "
          "blocks for a TCM clinical-decision-support knowledge base. Clinical "
          "validation through prospective randomised controlled trials is the "
          "obvious next step. The eleven analysis scripts, thirteen intermediate "
          "artefacts, six tables and eight figures are released as a single-"
          "command reproducibility package archived under a permanent DOI.",
          "MDPI_3.1_text")

    # ── Back matter ────────────────────────────────────────────────
    p = doc.add_paragraph(style="MDPI_6.2_back_matter")
    p.add_run("Author Contributions: ").bold = True
    p.add_run("Conceptualisation, J.L.; methodology, J.L.; software, J.L.; "
              "validation, J.L. and G.S.; formal analysis, J.L.; investigation, "
              "J.L.; resources, J.L.; data curation, J.L.; writing—original "
              "draft preparation, J.L.; writing—review and editing, G.S.; "
              "visualisation, J.L.; supervision, G.S.; project administration, "
              "G.S. All authors have read and agreed to the published version "
              "of the manuscript.")

    p = doc.add_paragraph(style="MDPI_6.2_back_matter")
    p.add_run("Funding: ").bold = True
    p.add_run("This research received no external funding.")

    p = doc.add_paragraph(style="MDPI_6.2_back_matter")
    p.add_run("Institutional Review Board Statement: ").bold = True
    p.add_run("The study did not involve new interventions in human subjects "
              "and did not access any institutional electronic medical record "
              "system. The prescription information used in this study was "
              "derived from publicly available literature, classical formulary "
              "databases and a research-curated, fully de-identified "
              "prescription archive. The analytical dataset retained only the "
              "two atomic fields prescription_id and herb and contained no "
              "information by which an individual patient, prescriber or "
              "institution could be identified. Under the Chinese Measures for "
              "Ethical Review of Life Science and Medical Research Involving "
              "Humans (2023) and equivalent international frameworks, this "
              "category of research is generally eligible for ethical review "
              "exemption; the corresponding author will provide a formal "
              "exemption confirmation prior to acceptance.")

    p = doc.add_paragraph(style="MDPI_6.2_back_matter")
    p.add_run("Informed Consent Statement: ").bold = True
    p.add_run("Because the present analysis used only de-identified records "
              "and publicly available sources, individual informed consent was "
              "not required.")

    p = doc.add_paragraph(style="MDPI_6.2_back_matter")
    p.add_run("Data Availability Statement: ").bold = True
    p.add_run("The two-field analytical dataset "
              "(anonymized_prescription_transactions.csv; 14,178 records / "
              "986 prescriptions) and the 24-rule herb-name normalisation "
              "ontology (herb_normalization_map.csv) are openly available "
              "together with the eleven analysis scripts and the manuscript "
              "source at https://github.com/ourHua/insomnia-tcm-mining. Each "
              "tagged release is permanently archived on Zenodo: the v1.0.2 "
              "archive carries the DOI 10.5281/zenodo.20250578. By design the "
              "released files do not enable re-identification of any patient, "
              "prescriber or institution. Aggregated source-class summaries "
              "are provided in Supplementary Table S1; the corresponding "
              "author will respond to methodological requests under appropriate "
              "data-use agreements.")

    p = doc.add_paragraph(style="MDPI_6.2_back_matter")
    p.add_run("Acknowledgments: ").bold = True
    p.add_run("The authors thank the developers of the open-source libraries "
              "on which this pipeline relies (pandas, NumPy, SciPy, mlxtend, "
              "NetworkX, python-louvain, leidenalg, scikit-learn, matplotlib).")

    p = doc.add_paragraph(style="MDPI_6.2_back_matter")
    p.add_run("Conflicts of Interest: ").bold = True
    p.add_run("The authors declare no conflict of interest.")

    # ── References ─────────────────────────────────────────────────
    _para(doc, "References", "MDPI_2.1_heading1")
    REFS = [
        "Benjafield, A.V.; Sert Kuniyoshi, F.H.; Malhotra, A.; Martin, J.L.; Morin, C.M.; Maurer, L.F.; Cistulli, P.A.; Pépin, J.L.; Wickwire, E.M.; medXcloud group. Estimation of the global prevalence and burden of insomnia: A systematic literature review-based analysis. Sleep Med. Rev. 2025, 82, 102121. https://doi.org/10.1016/j.smrv.2025.102121.",
        "Riemann, D.; Espie, C.A.; Altena, E.; et al. The European Insomnia Guideline: An update on the diagnosis and treatment of insomnia 2023. J. Sleep Res. 2023, 32, e14035. https://doi.org/10.1111/jsr.14035.",
        "Edinger, J.D.; Arnedt, J.T.; Bertisch, S.M.; et al. Behavioral and psychological treatments for chronic insomnia disorder in adults: An American Academy of Sleep Medicine clinical practice guideline. J. Clin. Sleep Med. 2021, 17, 255–262. https://doi.org/10.5664/jcsm.8986.",
        "Furukawa, Y.; Sakata, M.; Furukawa, T.A.; Efthimiou, O.; Perlis, M. Initial treatment choices for long-term remission of chronic insomnia disorder in adults: A systematic review and network meta-analysis. Psychiatry Clin. Neurosci. 2024, 78, 646–653. https://doi.org/10.1111/pcn.13730.",
        "Zhang, Y.; Ren, R.; Yang, L.; et al. Comparative efficacy and acceptability of psychotherapies, pharmacotherapies, and their combination for the treatment of adult insomnia: A systematic review and network meta-analysis. Sleep Med. Rev. 2022, 65, 101687. https://doi.org/10.1016/j.smrv.2022.101687.",
        "Ye, Z.; Lai, H.; Ning, J.; et al. Traditional Chinese medicine for insomnia: Recommendation mapping of the global clinical guidelines. J. Ethnopharmacol. 2024, 322, 117601. https://doi.org/10.1016/j.jep.2023.117601.",
        "Ma, N.; Pan, B.; Yang, S.; et al. Comparative efficacy and safety of Chinese patent medicines for primary insomnia: A systematic review and network meta-analysis of 109 randomized trials. J. Ethnopharmacol. 2025, 340, 119254. https://doi.org/10.1016/j.jep.2024.119254.",
        "Liu, X.X.; Ma, Y.Q.; Wang, Y.G.; et al. Suanzaoren decoction for the treatment of chronic insomnia: A systematic review and meta-analysis. Eur. Rev. Med. Pharmacol. Sci. 2022, 26, 8523–8533. https://doi.org/10.26355/eurrev_202211_30388.",
        "Tang, Y.; Li, Z.; Yang, D.; et al. Research of insomnia on traditional Chinese medicine diagnosis and treatment based on machine learning. Chin. Med. 2021, 16, 2. https://doi.org/10.1186/s13020-020-00409-8.",
        "Zhang, C.; Yang, X.; Ye, J.; et al. Mapping the research landscape of traditional Chinese medicine in insomnia management: A bibliometric study (2005–2024). Front. Neurol. 2025, 16, 1614948. https://doi.org/10.3389/fneur.2025.1614948.",
        "Agrawal, R.; Imieliński, T.; Swami, A. Mining association rules between sets of items in large databases. In Proceedings of the 1993 ACM SIGMOD International Conference on Management of Data, Washington, DC, USA, 26–28 May 1993; pp. 207–216. https://doi.org/10.1145/170035.170072.",
        "Yang, J.; Wang, S.; Zhang, Z.; et al. Analysis of medication rule of traditional Chinese medicine in treating depression based on data mining. Heliyon 2024, 10, e39245. https://doi.org/10.1016/j.heliyon.2024.e39245.",
        "Blondel, V.D.; Guillaume, J.L.; Lambiotte, R.; Lefebvre, E. Fast unfolding of communities in large networks. J. Stat. Mech. Theory Exp. 2008, 2008, P10008. https://doi.org/10.1088/1742-5468/2008/10/P10008.",
        "Traag, V.A.; Waltman, L.; van Eck, N.J. From Louvain to Leiden: Guaranteeing well-connected communities. Sci. Rep. 2019, 9, 5233. https://doi.org/10.1038/s41598-019-41695-z.",
        "Hu, F.; Li, L.; Huang, X.; Yan, X.; Huang, P. Symptom distribution regularity of insomnia: Network and spectral clustering analysis. JMIR Med. Inform. 2020, 8, e16749. https://doi.org/10.2196/16749.",
        "Yu, H.; Choi, K.; Kim, J.Y.; Yoo, S. Multi-level association rule mining and network pharmacology to identify the polypharmacological effects of herbal materials and compounds in traditional medicine. Brief. Bioinform. 2025, 26, bbaf328. https://doi.org/10.1093/bib/bbaf328.",
        "Wen, Z.; Dong, Y.; Peng, L.; Zhang, L.; Yan, J. PRDAGE: A prescription recommendation framework for traditional Chinese medicine based on data augmentation and multi-graph embedding. PeerJ Comput. Sci. 2025, 11, e2974. https://doi.org/10.7717/peerj-cs.2974.",
        "Wang, S.; Zhao, Y.; Hu, X. Exploring the mechanism of Suanzaoren decoction in treatment of insomnia based on network pharmacology and molecular docking. Front. Pharmacol. 2023, 14, 1145532. https://doi.org/10.3389/fphar.2023.1145532.",
        "Chu, Y.; Zhang, Y.; Liu, J.; Du, C.; Yan, Y. An integrated liver, hippocampus and serum metabolomics based on UPLC-Q-TOF-MS revealed the therapeutical mechanism of Ziziphi Spinosae Semen in p-chlorophenylalanine-induced insomnia rats. Biomed. Chromatogr. 2024, 38, e5796. https://doi.org/10.1002/bmc.5796.",
        "Yan, Y.; Zhao, N.; Liu, J.; et al. Ziziphi Spinosae Semen flavonoid ameliorates hypothalamic metabolism and modulates gut microbiota in chronic restraint stress-induced anxiety-like behavior in mice. Foods 2025, 14, 828. https://doi.org/10.3390/foods14050828.",
        "Xiao, F.; Shao, S.; Zhang, H.; et al. Neuroprotective effect of Ziziphi Spinosae Semen on rats with p-chlorophenylalanine-induced insomnia via activation of GABA-A receptor. Front. Pharmacol. 2022, 13, 965308. https://doi.org/10.3389/fphar.2022.965308.",
        "Sun, M.; Li, M.; Cui, X.; et al. Terpenoids derived from Semen Ziziphi Spinosae oil enhance sleep by modulating neurotransmitter signaling in mice. Heliyon 2024, 10, e26979. https://doi.org/10.1016/j.heliyon.2024.e26979.",
        "Jiang, N.; Wei, S.; Zhang, Y.; et al. Protective effects and mechanism of Radix Polygalae against neurological diseases as well as effective substance. Front. Psychiatry 2021, 12, 688703. https://doi.org/10.3389/fpsyt.2021.688703.",
        "Luo, H.; Sun, S.J.; Wang, Y.; Wang, Y.L. Revealing the sedative-hypnotic effect of the extracts of herb pair Semen Ziziphi spinosae and Radix Polygalae and related mechanisms through experiments and metabolomics approach. BMC Complement. Med. Ther. 2020, 20, 206. https://doi.org/10.1186/s12906-020-03000-8.",
        "Liu, Z.; Wang, Q.; Fan, X.; et al. Os Draconis-derived nanoparticles improve insomnia symptoms by activating calcium-dependent 5-HT release and the vagal-NTS pathway. Int. J. Nanomedicine 2025, 20, 14329–14341. https://doi.org/10.2147/IJN.S553405.",
        "Fan, J.; Li, Y.; Wang, G.; Li, Q.; Fang, M. Identification and verification of the anti-insomnia compounds from herbal pair Albiziae Cortex and Polygoni Multiflori Caulis using immobilized 5-HT1A receptor chromatography. Biomed. Chromatogr. 2026, 40, e70278. https://doi.org/10.1002/bmc.70278.",
        "Varinthra, P.; Anwar, S.N.M.N.; Shih, S.-C.; Liu, I.Y. The role of the GABAergic system on insomnia. Tzu Chi Med. J. 2024, 36, 103–109. https://doi.org/10.4103/tcmj.tcmj_243_23.",
        "Liu, X.; Sun, P.; Bao, X.; Cao, Y.; Wang, L.; Wang, Q. Potential mechanisms of traditional Chinese medicine in treating insomnia: A network pharmacology, GEO validation, and molecular-docking study. Medicine 2024, 103, e38052. https://doi.org/10.1097/MD.0000000000038052.",
        "Chung, M.C.; Su, L.J.; Chen, C.L.; Wu, L.C. AI-assisted literature exploration of innovative Chinese medicine formulas. Front. Pharmacol. 2024, 15, 1347882. https://doi.org/10.3389/fphar.2024.1347882.",
        "Zhao, W.; Lu, W.; Li, Z.; et al. TCM herbal prescription recommendation model based on multi-graph convolutional network. J. Ethnopharmacol. 2022, 297, 115109. https://doi.org/10.1016/j.jep.2022.115109.",
        "Han, X.; Xie, X.; Zhao, R.; et al. Calculating the similarity between prescriptions to find their new indications based on graph neural network. Chin. Med. 2024, 19, 124. https://doi.org/10.1186/s13020-024-00994-y.",
        "Zhao, F.-Y.; Xu, P.; Kennedy, G.A.; et al. Commercial Chinese polyherbal preparation Zao Ren An Shen prescription for primary insomnia: A systematic review with meta-analysis and trial sequential analysis. Front. Pharmacol. 2024, 15, 1376637. https://doi.org/10.3389/fphar.2024.1376637.",
        "Yang, M.; Wang, H.; Zhang, Y.L.; et al. The herbal medicine Suanzaoren (Ziziphi Spinosae Semen) for sleep quality improvements: A systematic review and meta-analysis. Integr. Cancer Ther. 2023, 22, 15347354231162080. https://doi.org/10.1177/15347354231162080.",
        "Wang, M.; Wang, G.; Zhao, M.; et al. Jujuboside A in ameliorating insomnia in mice via GABAergic modulation of the PVT. J. Ethnopharmacol. 2025, 349, 119939. https://doi.org/10.1016/j.jep.2025.119939.",
        "Bian, Z.; Zhang, W.; Feng, Z.; et al. Ziziphi Spinosae semen extract ameliorates insomnia by regulating hypothalamic GABA/Glu balance and gut microbiota Lactobacillus johnsonii. J. Funct. Foods 2025, 129, 106911. https://doi.org/10.1016/j.jff.2025.106911.",
        "Hagberg, A.A.; Schult, D.A.; Swart, P.J. Exploring network structure, dynamics, and function using NetworkX. In Proceedings of the 7th Python in Science Conference (SciPy 2008), Pasadena, CA, USA, 19–24 August 2008; pp. 11–15.",
        "Newman, M.E.J. Modularity and community structure in networks. Proc. Natl. Acad. Sci. USA 2006, 103, 8577–8582. https://doi.org/10.1073/pnas.0601602103.",
        "Watts, D.J.; Strogatz, S.H. Collective dynamics of 'small-world' networks. Nature 1998, 393, 440–442. https://doi.org/10.1038/30918.",
        "Freeman, L.C. Centrality in social networks: Conceptual clarification. Soc. Networks 1978, 1, 215–239. https://doi.org/10.1016/0378-8733(78)90021-7.",
        "Raschka, S. MLxtend: Providing machine learning and data science utilities and extensions to Python's scientific computing stack. J. Open Source Softw. 2018, 3, 638. https://doi.org/10.21105/joss.00638.",
    ]
    for r in REFS:
        p = doc.add_paragraph(style="MDPI_8.1_references")
        p.add_run(r)

    doc.save(str(OUT))
    print(f"WROTE: {OUT}  ({OUT.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    build()
